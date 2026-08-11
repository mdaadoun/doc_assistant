"""Unit tests for DOCX document parser and structural metadata extraction."""

from pathlib import Path

import docx
import pytest
from docx.enum.text import WD_BREAK

from core.exceptions import IngestionError
from ingestion.docx_parser import DOCXParser
from models.document import ParsedDocument


@pytest.fixture
def sample_docx(tmp_path: Path) -> Path:
    """Create a temporary multi-page DOCX file with headings, text, tables, and metadata."""
    docx_path = tmp_path / "sample.docx"
    doc = docx.Document()

    # Section / Core properties
    core = doc.core_properties
    core.title = "Helvetia Technical Architecture"
    core.author = "Engineering Team"
    core.subject = "System Specifications"

    # Heading and paragraph
    doc.add_heading("Section 1: Architecture Overview", level=1)
    doc.add_paragraph("This document outlines the corporate RAG architecture.")

    # Table
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Module"
    table.cell(0, 1).text = "Status"
    table.cell(1, 0).text = "DOCX Parser"
    table.cell(1, 1).text = "Complete"

    # Page break
    p_break = doc.add_paragraph("Section 2: Implementation Details")
    p_break.add_run().add_break(WD_BREAK.PAGE)
    doc.add_paragraph("Page two detailed technical guidelines.")

    doc.save(str(docx_path))
    return docx_path


def test_docx_parser_success(sample_docx: Path) -> None:
    """Verify DOCXParser extracts headings, tables, metadata, and handles pagination."""
    parser = DOCXParser()
    parsed: ParsedDocument = parser.parse(sample_docx)

    assert parsed.file_name == "sample.docx"
    assert parsed.source_format == "docx"
    assert parsed.doc_metadata.title == "Helvetia Technical Architecture"
    assert parsed.doc_metadata.author == "Engineering Team"
    assert parsed.doc_metadata.subject == "System Specifications"
    assert parsed.doc_metadata.total_pages >= 2
    assert len(parsed.pages) >= 2

    # Page 1 checks
    p1 = parsed.pages[0]
    assert p1.page_number == 1
    assert "# Section 1: Architecture Overview" in p1.text
    assert "Module | Status" in p1.text
    assert "DOCX Parser | Complete" in p1.text
    assert p1.metadata.table_count == 1
    assert p1.metadata.word_count > 0

    # Page 2 checks
    p2 = parsed.pages[1]
    assert p2.page_number == 2
    assert "Page two detailed technical guidelines." in p2.text


def test_docx_parser_nonexistent_file(tmp_path: Path) -> None:
    """Verify IngestionError is raised for non-existent file path."""
    parser = DOCXParser()
    missing_file = tmp_path / "missing.docx"
    with pytest.raises(IngestionError) as exc_info:
        parser.parse(missing_file)
    assert exc_info.value.code == "FILE_NOT_FOUND"


def test_docx_parser_empty_file(tmp_path: Path) -> None:
    """Verify IngestionError is raised for empty 0-byte DOCX file."""
    empty_file = tmp_path / "empty.docx"
    empty_file.write_bytes(b"")
    parser = DOCXParser()
    with pytest.raises(IngestionError) as exc_info:
        parser.parse(empty_file)
    assert exc_info.value.code == "EMPTY_FILE"


def test_docx_parser_corrupt_file(tmp_path: Path) -> None:
    """Verify IngestionError is raised for invalid corrupted DOCX file."""
    corrupt_file = tmp_path / "corrupt.docx"
    corrupt_file.write_bytes(b"Invalid zip archive docx data")
    parser = DOCXParser()
    with pytest.raises(IngestionError) as exc_info:
        parser.parse(corrupt_file)
    assert exc_info.value.code == "DOCX_PARSING_ERROR"
